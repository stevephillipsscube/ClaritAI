# build_project_info_from_layout.py
# Offline: extract "Project Information" fields from local Layout XML
# and (optionally) update a Python script's PROJECT_INFO block.

import re
import sys
import json
import argparse
from pathlib import Path
import xml.etree.ElementTree as ET
import csv
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

def _lname(tag: str) -> str:
    return tag.split('}', 1)[-1] if '}' in tag else tag

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """Insert a new paragraph *after* the given paragraph and return it."""
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p

def nice_label(api: str) -> str:
    name = api.split(".")[-1]
    name = name.replace("__c", "").replace("__", "_")
    parts = [w for w in name.split("_") if w]
    label = " ".join(("ID" if w.lower()=="id" else w.capitalize()) for w in parts)
    return label.replace("Url", "URL").replace("Zip", "ZIP")

TEMPLATE_PATH = Path(r"C:\Users\StevePhillips\Documents\Script Interface\ClaritAI\ClaritAI\TEMPLATE.docx")

def slug(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")

def find_candidate_layouts(layouts_dir: Path, sobject: str, app_type: str):
    # Prefer exact-ish matches; penalize "*Experience*" variants
    words = [w for w in re.split(r"[\s/_-]+", app_type.strip()) if w]
    files = list(layouts_dir.glob(f"{sobject}-*.layout-meta.xml"))
    ranked = []
    for p in files:
        n = p.name
        lname = n.lower()
        score = 0.0
        for w in words:
            if w.lower() in lname:
                score += 1.0
        # strong bonus if the literal phrase appears
        if app_type.lower() in lname:
            score += 1.5
        if "experience" in lname:
            score -= 0.25
        if score > 0:
            ranked.append((score, len(n), p))
    ranked.sort(key=lambda t: (-t[0], t[1]))
    return [p for _,__,p in ranked]

def parse_project_info_fields(layout_file: Path, debug=False):
    """
    Accepts section labels:
      - 'Project Information', 'Project Details', 'Project Info' (case-insensitive)
    Supports both shapes:
      A) layoutColumns/layoutItems/field
      B) layoutRows/layoutItems/layoutComponents/value   (value holds API name)
    Ignores namespace prefixes by matching on local-name.
    """
    ACCEPT = {"project information", "project details", "project info"}

    root = ET.parse(layout_file).getroot()
    all_fields = []

    # Walk all <layoutSections> regardless of namespace
    for sec in root.iter():
        if _lname(sec.tag) != "layoutSections":
            continue

        # read <label>
        label = ""
        for child in sec:
            if _lname(child.tag) == "label":
                label = (child.text or "").strip()
                break
        if label.lower() not in ACCEPT:
            continue

        # Shape A: <layoutColumns>/<layoutItems>/<field>
        for col in sec:
            if _lname(col.tag) != "layoutColumns":
                continue
            for it in col:
                if _lname(it.tag) != "layoutItems":
                    continue
                for fld in it:
                    if _lname(fld.tag) == "field" and (fld.text or "").strip():
                        all_fields.append(fld.text.strip())

        # Shape B: <layoutRows>/<layoutItems>/<layoutComponents>/<value>
        for row in sec:
            if _lname(row.tag) != "layoutRows":
                continue
            for it in row:
                if _lname(it.tag) != "layoutItems":
                    continue
                # skip emptySpace
                for es in it:
                    if _lname(es.tag) == "emptySpace" and (es.text or "").strip().lower() == "true":
                        break
                else:
                    for comp in it:
                        if _lname(comp.tag) != "layoutComponents":
                            continue
                        for val in comp:
                            if _lname(val.tag) == "value":
                                token = (val.text or "").strip()
                                # keep tokens that look like API names (no spaces/braces/dots)
                                if token and not any(ch in token for ch in (" ", ":", "/", "{", "}", ".")):
                                    all_fields.append(token)

    # de-dupe, preserve order
    seen, ordered = set(), []
    for f in all_fields:
        if f not in seen:
            seen.add(f)
            ordered.append(f)

    if debug:
        print(f"[debug] {layout_file.name}: found {len(ordered)} field(s)")
    return ordered


def write_merge_csv(sobject: str, app_type: str, fields: list[str], out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    csv_path = out_dir / f"project_info_tags_{slug(app_type)}.csv"
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Section", "Field Label", "API Name", "Merge Tag"])
        for api in fields:
            w.writerow(["Project Information", nice_label(api), api, f"«{sobject}.{api}»"])
    return csv_path

def format_project_info_block(fields: list[str]):
    items = [{"path": api, "label": nice_label(api)} for api in fields]
    js = json.dumps(items, indent=2).replace("true","True").replace("false","False").replace("null","None")
    return "PROJECT_INFO = " + js + "\n"

def update_script(script_path: Path, new_block: str):
    src = script_path.read_text(encoding="utf-8")

    # Prefer marker-based replace if present
    start_marker = "# --- PROJECT_INFO (AUTO-GENERATED START) ---"
    end_marker   = "# --- PROJECT_INFO (AUTO-GENERATED END) ---"
    if start_marker in src and end_marker in src:
        pattern = re.compile(re.escape(start_marker)+r".*?"+re.escape(end_marker), re.DOTALL)
        replacement = start_marker + "\n" + new_block + end_marker
        script_path.write_text(pattern.sub(replacement, src), encoding="utf-8")
        return True

    # Fallback: replace the first PROJECT_INFO = [ ... ] block
    pattern = re.compile(r"PROJECT_INFO\s*=\s*\[(?:.|\n)*?\]\s*", re.DOTALL)
    if not pattern.search(src):
        raise RuntimeError("Could not locate an existing PROJECT_INFO block in script.")
    script_path.write_text(pattern.sub(new_block, src, count=1), encoding="utf-8")
    return True

def inject_project_info_into_template(app_type: str, sobject: str, project_fields: list[str], out_dir: Path) -> Path:
    """
    Open TEMPLATE_PATH, find 'Project Information' or 'Project Details' heading,
    and inject lines like:  <Label>: «<SOBJECT>.<API>»
    Saves a new DOCX under /out and returns its path.
    """
    if not TEMPLATE_PATH.exists():
        raise SystemExit(f"Template not found: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    anchors = {"project information", "project details", "project info"}

    # Find the section heading paragraph
    anchor = None
    for p in doc.paragraphs:
        if (p.text or "").strip().lower() in anchors:
            anchor = p
            break

    # If not found, append a new heading at the end
    if anchor is None:
        anchor = doc.add_paragraph()
        r = anchor.add_run("Project Information")
        # (don’t force heading style; many templates have custom styles)

    # Build lines under the heading
    # Insert a blank spacer first (for clarity)
    cur = insert_paragraph_after(anchor, "")

    for api in project_fields:
        label = nice_label(api)
        line = f"{label}: «{sobject}.{api}»"
        cur = insert_paragraph_after(cur, line)

    # Save
    out_dir.mkdir(parents=True, exist_ok=True)
    safe = "".join(c if c.isalnum() or c in ("-","_") else "-" for c in app_type.strip().lower())
    out_path = out_dir / f"template_with_project_info_{safe}.docx"
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description="Extract 'Project Information' fields from a local Layout XML (offline).")
    ap.add_argument("application_type", help="e.g. 'Development Name Change'")
    ap.add_argument("--layouts", default="force-app/main/default/layouts", help="Layouts folder")
    ap.add_argument("--sobject", default="MUSW__Application2__c", help="sObject, e.g., MUSW__Application2__c")
    ap.add_argument("--out", default="out", help="Output folder for CSV")
    ap.add_argument("--update-script", default="", help="Python file whose PROJECT_INFO block should be replaced")
    ap.add_argument("--debug", action="store_true")
    args = ap.parse_args()

    layouts_dir = Path(args.layouts)
    if not layouts_dir.exists():
        print(f"✖ Layouts dir not found: {layouts_dir}")
        sys.exit(2)

    # Pick the best matching layout file
    candidates = find_candidate_layouts(layouts_dir, args.sobject, args.application_type)
    if not candidates:
        print(f"✖ No matching layouts for '{args.application_type}' under {layouts_dir}")
        sys.exit(3)

    target = candidates[0]
    print(f"→ Using layout: {target.name}")

    fields = parse_project_info_fields(target, debug=args.debug)
    if not fields:
        print("✖ No 'Project Information' section or fields found in the selected layout.")
        sys.exit(4)

    print(f"✓ Found {len(fields)} Project Information field(s):")
    for f in fields:
        print("   -", f)

    csv_path = write_merge_csv(args.sobject, args.application_type, fields, Path(args.out))
    print(f"✓ Wrote merge tags CSV: {csv_path}")

    if args.update_script:
        block = format_project_info_block(fields)
        try:
            update_script(Path(args.update_script), block)
            print(f"✓ Updated PROJECT_INFO in: {args.update_script}")
        except Exception as e:
            print(f"✖ Could not update script: {e}")
            sys.exit(5)

if __name__ == "__main__":
    main()
